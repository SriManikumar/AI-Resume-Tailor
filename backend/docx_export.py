from docx import Document

def markdown_to_docx(markdown_text: str) -> bytes:
    doc = Document()

    for line in markdown_text.split("\n"):
        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        # Bullet points
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.replace("- ", ""))
        else:
            doc.add_paragraph(line)

    # Save to memory
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
